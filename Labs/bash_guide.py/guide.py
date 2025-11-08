from docx import Document

doc = Document()
doc.add_heading('Step-by-Step Bash Scripting Learning Strategy', 0)

doc.add_paragraph("Bash scripting can feel overwhelming at first because it combines shell commands with programming logic, but with a clear step-by-step strategy, you can make steady progress. This guide provides a structured and simple approach to help you learn Bash scripting more effectively.")

steps = [
    {
        "title": "Step 1: Get Comfortable with the Linux Command Line",
        "content": [
            "Learn:",
            "- Navigating directories: cd, ls, pwd",
            "- File operations: cp, mv, rm, touch, mkdir",
            "- Viewing files: cat, less, head, tail",
            "- Permissions: chmod, chown",
            "- Pipes and redirects: |, >, >>, <, 2>",
            "Goal: Use the terminal daily for basic tasks."
        ]
    },
    {
        "title": "Step 2: Learn Basic Script Structure",
        "content": [
            "Learn:",
            "- Creating a script: nano myscript.sh",
            "- Adding the shebang: #!/bin/bash",
            "- Making it executable: chmod +x myscript.sh",
            "- Running it: ./myscript.sh",
            "Example:",
            "#!/bin/bash\necho \"Hello, world!\"",
            "Practice: Write 5 simple scripts that echo something or print the current date."
        ]
    },
    {
        "title": "Step 3: Variables and User Input",
        "content": [
            "Learn:",
            "- Defining variables: name=\"John\"",
            "- Using variables: echo \"Hello $name\"",
            "- Reading input: read -p \"Enter your name: \" name",
            "Try: Write a script that greets the user by name."
        ]
    },
    {
        "title": "Step 4: Conditions and If Statements",
        "content": [
            "Learn:",
            "- if statements:",
            "if [ $age -ge 18 ]; then\n  echo \"You're an adult.\"\nelse\n  echo \"You're underage.\"\nfi",
            "- Use comparison operators: -eq, -ne, -lt, -le, -gt, -ge, ==, !=",
            "Exercise: Write a script that checks if a file exists."
        ]
    },
    {
        "title": "Step 5: Loops (for, while)",
        "content": [
            "Learn:",
            "- for loop:",
            "for i in 1 2 3; do\n  echo \"Number $i\"\ndone",
            "- while loop:",
            "while [ $count -lt 5 ]; do\n  echo \"Count is $count\"\n  ((count++))\ndone",
            "Try: Write a loop that prints numbers 1 to 10."
        ]
    },
    {
        "title": "Step 6: Functions",
        "content": [
            "Learn:",
            "greet() {\n  echo \"Hello, $1!\"\n}\ngreet \"Alice\"",
            "Use case: Create a function that checks disk space."
        ]
    },
    {
        "title": "Step 7: Real-World Projects",
        "content": [
            "Apply what you've learned by building scripts that:",
            "- Backup files or folders",
            "- Monitor disk usage",
            "- Rename a bunch of files",
            "- Automate updates or installations"
        ]
    },
    {
        "title": "Bonus: Recommended Learning Resources",
        "content": [
            "- Free: https://explainshell.com – explains any bash command line",
            "- Interactive: https://www.learnshell.org/",
            "- Book: The Linux Command Line by William Shotts (free online)"
        ]
    },
    {
        "title": "Weekly Bash Plan",
        "content": [
            "Week 1: Terminal basics – Use Bash daily",
            "Week 2: Script syntax + echo + run – Write small scripts",
            "Week 3: Variables and input – Scripts with user interaction",
            "Week 4: If/else conditions – Logic-based scripts",
            "Week 5: Loops – Repetitive tasks automated",
            "Week 6: Functions – Modular and clean scripts",
            "Week 7+: Projects + daily practice – Real-world automation"
        ]
    }
]

for step in steps:
    doc.add_heading(step["title"], level=2)
    for line in step["content"]:
        doc.add_paragraph(line, style='Normal')

doc.save("Bash_Scripting_Learning_Strategy.docx")
